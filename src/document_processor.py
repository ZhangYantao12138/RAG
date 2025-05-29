"""
文档处理模块
"""
import os
import re
from typing import List, Optional
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""]
        )
    
    def validate_document(self, file_path: str) -> bool:
        """验证文档是否存在且格式正确"""
        try:
            if not os.path.exists(file_path):
                logger.error(f"文档文件不存在: {file_path}")
                return False
            
            if not file_path.lower().endswith('.docx'):
                logger.error(f"不支持的文档格式，仅支持.docx格式: {file_path}")
                return False
            
            # 尝试打开文档验证格式
            Document(file_path)
            logger.info(f"文档验证成功: {file_path}")
            return True
            
        except Exception as e:
            logger.error(f"文档验证失败: {e}")
            return False
    
    def load_document(self, file_path: str) -> str:
        """加载Word文档并提取文本内容"""
        try:
            if not self.validate_document(file_path):
                raise ValueError(f"文档验证失败: {file_path}")
            
            logger.info(f"开始加载文档: {file_path}")
            doc = Document(file_path)
            
            # 提取所有段落文本
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # 跳过空段落
                    text_content.append(paragraph.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("文档内容为空")
            
            logger.info(f"文档加载成功，提取了 {len(full_text)} 个字符")
            return full_text
            
        except Exception as e:
            logger.error(f"文档加载失败: {e}")
            raise
    
    def clean_text(self, text: str) -> str:
        """清洗文本内容"""
        if not text:
            return ""
        
        logger.info("开始清洗文本内容")
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符但保留中文标点
        text = re.sub(r'[^\u4e00-\u9fff\w\s，。！？：；""''（）【】、]', '', text)
        
        # 移除多余的换行符
        text = re.sub(r'\n+', '\n', text)
        
        # 移除前后空白
        text = text.strip()
        
        logger.info(f"文本清洗完成，最终长度: {len(text)}")
        return text
    
    def split_text(self, text: str) -> List[str]:
        """将文本分割成块"""
        if not text:
            return []
        
        logger.info(f"开始分割文本，原始长度: {len(text)}")
        
        # 使用LangChain的文本分割器
        chunks = self.text_splitter.split_text(text)
        
        # 过滤空块和过短的块
        filtered_chunks = []
        for chunk in chunks:
            chunk = chunk.strip()
            if len(chunk) >= 50:  # 最小块长度
                filtered_chunks.append(chunk)
        
        logger.info(f"文本分割完成，共生成 {len(filtered_chunks)} 个文本块")
        return filtered_chunks
    
    def process_document(self, file_path: str) -> List[str]:
        """完整的文档处理流程"""
        try:
            logger.info(f"开始处理文档: {file_path}")
            
            # 1. 加载文档
            raw_text = self.load_document(file_path)
            
            # 2. 清洗文本
            cleaned_text = self.clean_text(raw_text)
            
            # 3. 分割文本
            text_chunks = self.split_text(cleaned_text)
            
            logger.info(f"文档处理完成，共生成 {len(text_chunks)} 个文本块")
            return text_chunks
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            raise
    
    def get_document_stats(self, file_path: str) -> dict:
        """获取文档统计信息"""
        try:
            raw_text = self.load_document(file_path)
            cleaned_text = self.clean_text(raw_text)
            chunks = self.split_text(cleaned_text)
            
            return {
                "file_path": file_path,
                "file_size_mb": round(os.path.getsize(file_path) / 1024 / 1024, 2),
                "raw_text_length": len(raw_text),
                "cleaned_text_length": len(cleaned_text),
                "total_chunks": len(chunks),
                "avg_chunk_size": round(sum(len(chunk) for chunk in chunks) / len(chunks), 2) if chunks else 0
            }
        except Exception as e:
            logger.error(f"获取文档统计信息失败: {e}")
            return {} 