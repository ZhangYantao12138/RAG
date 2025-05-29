"""
文档格式转换模块
用于将Word文档转换为Markdown格式
"""
import os
import re
import logging
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from fuzzywuzzy import fuzz
from fuzzywuzzy import process

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    encoding='utf-8'  # 添加编码设置
)
logger = logging.getLogger(__name__)

class FormatConverter:
    """文档格式转换器"""
    
    def __init__(self):
        # 场景标记正则表达式
        self.scene_pattern = re.compile(r'^场景\s*\d+[：:].*$|^第[一二三四五六七八九十百千万]+场[：:].*$')
        # 对话标记正则表达式
        self.dialogue_pattern = re.compile(r'^[^：:]+[：:].*$')
        # 动作描述正则表达式
        self.action_pattern = re.compile(r'^[（(].*[)）]$')
        
        # 正确的人名列表
        self.correct_names = [
            '程聿怀',
            '羌青瓷',
            '程走柳',
            '蒋伯驾',
            '黛利拉',
            '阿奇',
            '以撒',
            # 可以在这里添加更多正确的人名
        ]
        
        # 模糊匹配的阈值（0-100）
        self.fuzzy_threshold = 60
        
        # # 已知的OCR错误修正字典（用于完全匹配的情况）
        # self.ocr_corrections = {
        #     '程丰怀': '程聿怀',
        #     # 可以在这里添加更多的错误修正
        # }
    
    def find_similar_name(self, text: str) -> Optional[str]:
        """使用模糊匹配查找最相似的正确人名"""
        # 如果文本完全匹配某个正确人名，直接返回
        if text in self.correct_names:
            return text
            
        # 使用模糊匹配找到最相似的人名
        result = process.extractOne(
            text,
            self.correct_names,
            scorer=fuzz.ratio,
            score_cutoff=self.fuzzy_threshold
        )
        
        if result:
            return result[0]
        return None

    def convert_to_markdown(self, file_path: str) -> str:
        """将Word文档转换为Markdown格式"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            if not file_path.lower().endswith('.docx'):
                raise ValueError(f"不支持的文件格式，仅支持.docx格式: {file_path}")
            
            logger.info(f"开始转换文档: {file_path}")
            
            # 加载Word文档
            doc = Document(file_path)
            
            # 转换内容
            markdown_lines = []
            
            # 处理段落
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                # 处理场景标记
                if self.scene_pattern.match(para.text):
                    markdown_lines.append(f"\n## {para.text}\n")
                    continue
                
                # 处理对话
                if self.dialogue_pattern.match(para.text):
                    # 分割说话人和内容
                    speaker, content = para.text.split('：', 1) if '：' in para.text else para.text.split(':', 1)
                    markdown_lines.append(f"**{speaker}**：{content}\n")
                    continue
                
                # 处理动作描述
                if self.action_pattern.match(para.text):
                    markdown_lines.append(f"*{para.text}*\n")
                    continue
                
                # 处理普通段落
                markdown_lines.append(f"{para.text}\n")
            
            # 处理表格
            for table in doc.tables:
                markdown_lines.append("\n")
                # 处理表头
                header = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_lines.append("| " + " | ".join(header) + " |")
                markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                
                # 处理表格内容
                for row in table.rows[1:]:
                    row_text = [cell.text.strip() for cell in row.cells]
                    markdown_lines.append("| " + " | ".join(row_text) + " |")
                markdown_lines.append("\n")
            
            # 合并所有行
            markdown_text = "".join(markdown_lines)
            
            # 清理多余的空行
            markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
            
            logger.info("文档转换完成")
            return markdown_text
            
        except Exception as e:
            logger.error(f"文档转换失败: {e}")
            raise
    
    def save_markdown(self, markdown_text: str, output_path: str) -> bool:
        """保存Markdown文件"""
        try:
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 保存文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_text)
            
            logger.info(f"Markdown文件已保存: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"保存Markdown文件失败: {e}")
            return False
    
    def process_document(self, input_path: str, output_path: Optional[str] = None) -> str:
        """处理文档转换的完整流程"""
        try:
            # 如果没有指定输出路径，则使用默认路径
            if not output_path:
                output_path = os.path.splitext(input_path)[0] + '.md'
            
            # 1. 先进行格式转换
            markdown_text = self.convert_to_markdown(input_path)
            
            # 2. 再进行文本处理
            processed_text = self.process_text(markdown_text)
            
            # 3. 保存处理后的文件
            if self.save_markdown(processed_text, output_path):
                return output_path
            else:
                raise Exception("保存Markdown文件失败")
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            raise

    def process_text(self, text: str) -> str:
        """处理文本内容"""
        # 删除所有半角空格和全角空格
        text = text.replace(' ', '').replace('　', '')
        
        # 删除所有*符号
        text = text.replace('*', '')
            
        # 使用正则表达式找出可能的人名
        # 这里假设人名后面跟着冒号或出现在对话中
        name_pattern = re.compile(r'([^：:]+)[：:]|（([^）]+)）')
        
        def replace_name(match):
            name = match.group(1) or match.group(2)
            if name:
                similar_name = self.find_similar_name(name)
                if similar_name and similar_name != name:
                    logger.info(f"修正人名: {name} -> {similar_name}")
                    return similar_name + match.group(0)[len(name):]
            return match.group(0)
        
        # 应用模糊匹配修正
        text = name_pattern.sub(replace_name, text)

        # 英文标点转中文标点
        en2zh_punct = {
            ',': '，', '.': '。', '?': '？', '!': '！', ':': '：', ';': '；',
            '"': '」', '\'': '』', '(': '（', ')': '）', '[': '【', ']': '】', '{': '｛', '}': '｝',
            '」': '"',  # 添加」到"的转换
        }
        for en, zh in en2zh_punct.items():
            text = text.replace(en, zh)

        # 只保留句末标点或引号后的换行符，其余换行符全部去掉
        keep_punct = r'([。！？；…」』’”\'\"])'
        text = re.sub(keep_punct + r'\n', r'\1<KEEP_NL>', text)
        text = text.replace('\n', '')
        text = text.replace('<KEEP_NL>', '\n')

        # 处理段落
        paragraphs = text.split('\n\n')
        processed_paragraphs = []
        for para in paragraphs:
            if para.strip():
                # 处理粗体文本
                para = re.sub(r'\*\*(.*?)\*\*', r'**\1**', para)
                processed_paragraphs.append(para)
        
        # 重新组合文本，确保段落之间只有一个空行
        return '\n\n'.join(processed_paragraphs)

class SmartTextSplitter:
    def __init__(self):
        self.boundary_detector = SemanticBoundaryDetector()
        self.context_analyzer = ContextAnalyzer(self.boundary_detector)
        
        # 切片大小配置
        self.min_chunk_size = 150    # 最小150字符
        self.max_chunk_size = 1000   # 最大1000字符
        self.overlap_size = 50       # 上下文重叠大小（字符数）
    
    def split_text(self, text):
        """智能文本切片，保持上下文重叠"""
        # 检测语义边界
        boundaries = self.boundary_detector.detect_boundaries(text)
        
        # 分析上下文关联
        context_scores = self.context_analyzer.analyze_context(text)
        
        # 将文本分割成句子
        sentences = self.split_into_sentences(text)
        
        # 根据语义边界和上下文关联进行切片
        chunks = []
        current_chunk = []
        current_size = 0
        overlap_buffer = []  # 用于存储重叠部分的句子
        
        for i, sentence in enumerate(sentences):
            current_chunk.append(sentence)
            current_size += len(sentence)
            
            # 检查是否需要切分
            if (i in boundaries or  # 语义边界
                current_size >= self.max_chunk_size or  # 大小限制
                (current_size >= self.min_chunk_size and context_scores[i] < 0.5)):  # 上下文关联度低
                
                if current_chunk:
                    # 保存当前块
                    chunks.append(''.join(current_chunk))
                    
                    # 准备重叠部分
                    overlap_buffer = []
                    overlap_size = 0
                    
                    # 从当前块的末尾开始，收集重叠部分
                    for sent in reversed(current_chunk):
                        if overlap_size + len(sent) <= self.overlap_size:
                            overlap_buffer.insert(0, sent)
                            overlap_size += len(sent)
                        else:
                            break
                    
                    # 重置当前块，但保留重叠部分
                    current_chunk = overlap_buffer.copy()
                    current_size = overlap_size
        
        # 添加最后一个块
        if current_chunk:
            chunks.append(''.join(current_chunk))
        
        return chunks
    
    def split_into_sentences(self, text):
        """将文本分割成句子"""
        # 使用标点符号作为句子分隔符
        sentence_endings = ['。', '！', '？', '…', '；', '\n']
        sentences = []
        current = []
        
        for char in text:
            current.append(char)
            if char in sentence_endings:
                sentences.append(''.join(current))
                current = []
        
        if current:  # 处理最后一个句子
            sentences.append(''.join(current))
            
        return sentences 